#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DocForge Python 文档处理器
功能：
1. 从 DOCX 模板提取样式规则
2. 将 Markdown 转换为带模板样式的 DOCX
"""

import sys
import json
import re
import os
import zipfile
from pathlib import Path
from urllib.parse import unquote, urlparse
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field, asdict
import xml.etree.ElementTree as ET

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
except ImportError:
    print("请先安装依赖: pip install python-docx lxml", file=sys.stderr)
    sys.exit(1)


# ─────────────────────────────────────────────
# 样式数据结构
# ─────────────────────────────────────────────

def default_style_rules() -> dict:
    return {
        "title": {
            "font": {"name": "黑体", "size": 22, "bold": True, "italic": False},
            "paragraph": {"alignment": "center", "space_before": 12, "space_after": 6, "line_spacing": 1.5}
        },
        "heading1": {
            "font": {"name": "黑体", "size": 16, "bold": True, "italic": False},
            "paragraph": {"alignment": "left", "space_before": 12, "space_after": 6, "line_spacing": 1.5}
        },
        "heading2": {
            "font": {"name": "楷体", "size": 14, "bold": True, "italic": False},
            "paragraph": {"alignment": "left", "space_before": 10, "space_after": 4, "line_spacing": 1.5}
        },
        "heading3": {
            "font": {"name": "宋体", "size": 12, "bold": True, "italic": False},
            "paragraph": {"alignment": "left", "space_before": 8, "space_after": 4, "line_spacing": 1.5}
        },
        "body": {
            "font": {"name": "宋体", "size": 12, "bold": False, "italic": False},
            "paragraph": {
                "alignment": "justify",
                "space_before": 0,
                "space_after": 4,
                "line_spacing": 1.5,
                "indent_first_line": 0.35
            }
        },
        "list": {
            "font": {"name": "宋体", "size": 12, "bold": False, "italic": False},
            "paragraph": {"alignment": "left", "space_before": 2, "space_after": 2, "line_spacing": 1.5}
        },
        "quote": {
            "font": {"name": "楷体", "size": 11, "bold": False, "italic": True},
            "paragraph": {"alignment": "left", "space_before": 4, "space_after": 4, "line_spacing": 1.5, "indent_left": 0.4}
        },
        "code": {
            "font": {"name": "Consolas", "size": 10, "bold": False, "italic": False},
            "paragraph": {"alignment": "left", "space_before": 6, "space_after": 6, "line_spacing": 1.2, "indent_left": 0.4}
        },
        "page_margin": {"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
        "page_size": {"width": 8.27, "height": 11.69, "orientation": "portrait"}
    }


# ─────────────────────────────────────────────
# 字体工具
# ─────────────────────────────────────────────

def set_run_font(run, font_name: str, size_pt: float, bold: bool = False, italic: bool = False, color_hex: str = None):
    """设置 run 的字体，正确处理中文字体"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic

    if color_hex:
        color_hex = color_hex.lstrip('#')
        if len(color_hex) == 6:
            run.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16)
            )

    # 设置字体名称（同时设置 ASCII 和东亚字体）
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()

    # 设置东亚字体（中文）
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_paragraph_format(para, style_cfg: dict):
    """设置段落格式"""
    pf = para.paragraph_format

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    align = style_cfg.get("alignment", "left")
    para.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    space_before = style_cfg.get("space_before", 0)
    space_after = style_cfg.get("space_after", 0)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

    line_spacing = style_cfg.get("line_spacing", 1.5)
    if line_spacing:
        pf.line_spacing = line_spacing          # float = 倍数行距
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    indent_first = style_cfg.get("indent_first_line", 0)
    if indent_first:
        pf.first_line_indent = Inches(indent_first)

    indent_left = style_cfg.get("indent_left", 0)
    if indent_left:
        pf.left_indent = Inches(indent_left)


def alignment_to_name(alignment) -> str:
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    return "left"


def length_to_pt(value) -> Optional[float]:
    return round(value.pt, 2) if value is not None else None


def length_to_inches(value) -> Optional[float]:
    return round(value.inches, 2) if value is not None else None


def get_run_font_name(run, fallback: str) -> str:
    if run.font.name:
        return run.font.name

    r_pr = run._r.rPr
    r_fonts = getattr(r_pr, "rFonts", None) if r_pr is not None else None
    if r_fonts is not None:
        east_asia = r_fonts.get(qn('w:eastAsia'))
        ascii_font = r_fonts.get(qn('w:ascii'))
        if east_asia:
            return east_asia
        if ascii_font:
            return ascii_font

    return fallback


def paragraph_to_rule(para, fallback_rule: dict) -> dict:
    """从实际段落直接格式中提取样式，优先使用 run 直接格式，再回退到段落样式。"""
    rule = json.loads(json.dumps(fallback_rule))
    style_font = getattr(para.style, "font", None)
    first_run = next((run for run in para.runs if run.text.strip()), None)

    if first_run:
        fallback_font = rule["font"].get("name", "宋体")
        if style_font is not None and style_font.name:
            fallback_font = style_font.name

        rule["font"]["name"] = get_run_font_name(first_run, fallback_font)

        size = first_run.font.size or (style_font.size if style_font is not None else None)
        if size is not None:
            rule["font"]["size"] = round(size.pt, 1)

        bold = first_run.font.bold
        if bold is None and style_font is not None:
            bold = style_font.bold
        if bold is not None:
            rule["font"]["bold"] = bool(bold)

        italic = first_run.font.italic
        if italic is None and style_font is not None:
            italic = style_font.italic
        if italic is not None:
            rule["font"]["italic"] = bool(italic)

    pf = para.paragraph_format
    if para.alignment is not None:
        rule["paragraph"]["alignment"] = alignment_to_name(para.alignment)

    space_before = length_to_pt(pf.space_before)
    space_after = length_to_pt(pf.space_after)
    first_line = length_to_inches(pf.first_line_indent)
    left_indent = length_to_inches(pf.left_indent)

    if space_before is not None:
        rule["paragraph"]["space_before"] = space_before
    if space_after is not None:
        rule["paragraph"]["space_after"] = space_after
    if first_line is not None:
        rule["paragraph"]["indent_first_line"] = first_line
    if left_indent is not None:
        rule["paragraph"]["indent_left"] = left_indent
    if isinstance(pf.line_spacing, (int, float)):
        rule["paragraph"]["line_spacing"] = round(float(pf.line_spacing), 2)

    return rule


def normalized_para_text(para) -> str:
    return re.sub(r"\s+", "", para.text or "")


def heading_level_from_text(text: str) -> Optional[str]:
    if re.match(r"^[一二三四五六七八九十]+[、.．]", text):
        return "heading1"
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return "heading2"
    if re.match(r"^\d+[、.．]", text):
        return "heading3"
    return None


def add_field(run, field_code: str, placeholder_text: str = '1') -> None:
    """在 run 中插入 Word 域代码。"""
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = field_code

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    text = OxmlElement('w:t')
    text.text = placeholder_text

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_page_number_footer(doc: Document) -> None:
    """添加页脚页码，保证 DOCX 看起来像正式交付件。"""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run = para.add_run('第 ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run2 = para.add_run()
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_field(run2, 'PAGE')
    run3 = para.add_run(' 页 / 共 ')
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run3 = para.add_run()
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_field(run3, 'NUMPAGES')
    run4 = para.add_run(' 页')
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def validate_docx_package(docx_path: str) -> List[str]:
    """轻量校验 DOCX 包结构，避免把半成品当成正式交付件。"""
    errors = []
    required_entries = [
        '[Content_Types].xml',
        '_rels/.rels',
        'word/document.xml',
        'word/_rels/document.xml.rels',
        'docProps/core.xml',
    ]

    if not os.path.exists(docx_path):
        return [f'文件不存在: {docx_path}']

    if os.path.getsize(docx_path) <= 0:
        return [f'文件为空: {docx_path}']

    try:
        with zipfile.ZipFile(docx_path, 'r') as archive:
            names = set(archive.namelist())
            for entry in required_entries:
                if entry not in names:
                    errors.append(f'缺少 DOCX 组件: {entry}')

            bad_file = archive.testzip()
            if bad_file:
                errors.append(f'DOCX 压缩包损坏: {bad_file}')

            if 'word/document.xml' in names:
                document_xml = archive.read('word/document.xml')
                try:
                    root = ET.fromstring(document_xml)
                    body = root.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
                    if body is None or len(list(body)) == 0:
                        errors.append('word/document.xml 缺少正文内容')
                except ET.ParseError as exc:
                    errors.append(f'word/document.xml 解析失败: {exc}')
    except zipfile.BadZipFile:
        errors.append('文件不是有效的 DOCX zip 包')
    except Exception as exc:
        errors.append(f'DOCX 校验失败: {exc}')

    return errors


# ─────────────────────────────────────────────
# Markdown 解析
# ─────────────────────────────────────────────

CALLOUT_LABELS = {
    "note": "提示",
    "info": "信息",
    "todo": "待办",
    "tip": "建议",
    "success": "完成",
    "question": "问题",
    "warning": "注意",
    "failure": "失败",
    "danger": "风险",
    "bug": "问题",
    "example": "示例",
    "quote": "引用",
}


def normalize_wikilink(raw: str) -> str:
    """将 Obsidian [[target|alias]] 链接转换为适合正式文档展示的文本。"""
    target = raw.strip()
    if '|' in target:
        _, alias = target.split('|', 1)
        return alias.strip() or target.strip()
    return target.replace('\\', '/').split('/')[-1].strip()


def parse_inline(text: str) -> List[Tuple[str, dict]]:
    """解析行内格式，返回 (text, attrs) 列表。兼容 Obsidian 常用链接与脚注语法。"""
    result: List[Tuple[str, dict]] = []
    pos = 0
    special_chars = ['[', '!', '`', '*', '_', '~']

    patterns = [
        (re.compile(r'\[\^([^\]]+)\]'), lambda m: (f"[{m.group(1)}]", {'superscript': True})),
        (re.compile(r'\[\[([^\]]+)\]\]'), lambda m: (normalize_wikilink(m.group(1)), {'internal_link': True})),
        (re.compile(r'!\[([^\]]*)\]\(([^)]+)\)'), lambda m: (m.group(1) or m.group(2), {'image_ref': m.group(2)})),
        (re.compile(r'\[([^\]]+)\]\(([^)]+)\)'), lambda m: (m.group(1), {'link': m.group(2)})),
        (re.compile(r'`([^`]+)`'), lambda m: (m.group(1), {'code': True})),
        (re.compile(r'\*\*\*([^*]+)\*\*\*'), lambda m: (m.group(1), {'bold': True, 'italic': True})),
        (re.compile(r'___([^_]+)___'), lambda m: (m.group(1), {'bold': True, 'italic': True})),
        (re.compile(r'\*\*([^*]+)\*\*'), lambda m: (m.group(1), {'bold': True})),
        (re.compile(r'__([^_]+)__'), lambda m: (m.group(1), {'bold': True})),
        (re.compile(r'\*([^*\s][^*]*?)\*'), lambda m: (m.group(1), {'italic': True})),
        (re.compile(r'_([^_\s][^_]*?)_'), lambda m: (m.group(1), {'italic': True})),
        (re.compile(r'~~(.+?)~~'), lambda m: (m.group(1), {'strike': True})),
    ]

    while pos < len(text):
        matched = False
        chunk = text[pos:]
        for pattern, mapper in patterns:
            match = pattern.match(chunk)
            if not match:
                continue
            mapped_text, attrs = mapper(match)
            if mapped_text:
                result.append((mapped_text, attrs))
            pos += match.end()
            matched = True
            break

        if matched:
            continue

        next_positions = [text.find(ch, pos + 1) for ch in special_chars]
        next_positions = [p for p in next_positions if p != -1]
        next_pos = min(next_positions) if next_positions else len(text)
        result.append((text[pos:next_pos], {}))
        pos = next_pos

    return result if result else [(text, {})]


def strip_obsidian_comments(line: str) -> str:
    """移除单行 Obsidian 注释片段。"""
    return re.sub(r'%%.*?%%', '', line).strip()


def parse_image_marker(stripped: str) -> Optional[dict]:
    """解析 Markdown 图片和 Obsidian 嵌入图片。"""
    md_image = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
    if md_image:
        return {'type': 'image', 'alt': md_image.group(1).strip(), 'src': md_image.group(2).strip()}

    obsidian_embed = re.match(r'^!\[\[([^\]]+)\]\]\s*$', stripped)
    if obsidian_embed:
        raw = obsidian_embed.group(1).strip()
        src, _, size_or_alt = raw.partition('|')
        return {'type': 'image', 'alt': size_or_alt.strip(), 'src': src.strip()}

    return None


def parse_markdown(markdown: str) -> List[dict]:
    """解析 Markdown 为元素列表，重点兼容 Obsidian 到正式 DOCX 的常见落差。"""
    elements: List[dict] = []
    footnotes: List[dict] = []
    lines = markdown.split('\n')
    i = 0
    in_code_block = False
    in_comment_block = False
    code_content: List[str] = []
    code_lang = ''
    paragraph_lines: List[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        content = ' '.join(part.strip() for part in paragraph_lines if part.strip()).strip()
        if content:
            elements.append({'type': 'body', 'content': content})
        paragraph_lines.clear()

    while i < len(lines):
        line = lines[i].rstrip('\r')
        stripped = line.strip()

        if i == 0 and stripped == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        if in_comment_block:
            if stripped.endswith('%%'):
                in_comment_block = False
            i += 1
            continue

        if stripped.startswith('%%'):
            if not stripped.endswith('%%') or stripped == '%%':
                in_comment_block = True
            i += 1
            continue

        stripped = strip_obsidian_comments(stripped)

        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                flush_paragraph()
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_content = []
            else:
                in_code_block = False
                elements.append({'type': 'code', 'content': '\n'.join(code_content), 'lang': code_lang})
                code_content = []
                code_lang = ''
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        footnote_match = re.match(r'^\[\^([^\]]+)\]:\s*(.*)$', stripped)
        if footnote_match:
            flush_paragraph()
            footnotes.append({'id': footnote_match.group(1), 'content': footnote_match.group(2).strip()})
            i += 1
            continue

        image_element = parse_image_marker(stripped)
        if image_element:
            flush_paragraph()
            elements.append(image_element)
            i += 1
            continue

        # 标题
        if stripped.startswith('#### '):
            flush_paragraph()
            elements.append({'type': 'heading4', 'content': stripped[5:]})
        elif stripped.startswith('### '):
            flush_paragraph()
            elements.append({'type': 'heading3', 'content': stripped[4:]})
        elif stripped.startswith('## '):
            flush_paragraph()
            elements.append({'type': 'heading2', 'content': stripped[3:]})
        elif stripped.startswith('# '):
            flush_paragraph()
            elements.append({'type': 'title', 'content': stripped[2:]})

        # 分隔线
        elif re.match(r'^[-*_]{3,}$', stripped):
            flush_paragraph()
            elements.append({'type': 'hr'})

        # 引用 / Obsidian callout
        elif stripped.startswith('>'):
            flush_paragraph()
            quote_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_line = re.sub(r'^>\s?', '', lines[i].strip())
                quote_lines.append(quote_line)
                i += 1

            callout = re.match(r'^\[!([A-Za-z]+)\][+-]?\s*(.*)$', quote_lines[0] if quote_lines else '')
            if callout:
                callout_type = callout.group(1).lower()
                title = callout.group(2).strip()
                content = ' '.join(line.strip() for line in quote_lines[1:] if line.strip()).strip()
                elements.append({
                    'type': 'callout',
                    'callout_type': callout_type,
                    'title': title or CALLOUT_LABELS.get(callout_type, callout_type),
                    'content': content,
                })
            else:
                content = ' '.join(line.strip() for line in quote_lines if line.strip()).strip()
                elements.append({'type': 'quote', 'content': content})
            continue

        # 无序列表 / 任务列表
        elif re.match(r'^\s*[-*+]\s', line):
            flush_paragraph()
            match = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
            indent = len((match.group(1) if match else '').replace('\t', '    '))
            content = match.group(2).strip() if match else re.sub(r'^[-*+]\s+', '', stripped)
            task = re.match(r'^\[([ xX])\]\s+(.*)$', content)
            task_state = None
            if task:
                task_state = 'done' if task.group(1).lower() == 'x' else 'todo'
                content = task.group(2).strip()
            elements.append({
                'type': 'list_item',
                'content': content,
                'ordered': False,
                'level': indent // 2,
                'task_state': task_state,
            })

        # 有序列表
        elif re.match(r'^\s*\d+\.\s', line):
            flush_paragraph()
            match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
            indent = len((match.group(1) if match else '').replace('\t', '    '))
            content = match.group(3).strip() if match else re.sub(r'^\d+\.\s+', '', stripped)
            number = int(match.group(2)) if match else None
            elements.append({'type': 'list_item', 'content': content, 'ordered': True, 'number': number, 'level': indent // 2})

        # 表格
        elif '|' in stripped and stripped.startswith('|'):
            flush_paragraph()
            table_rows, table_alignments, consumed_count = _parse_table_rows(lines, i)
            if table_rows:
                elements.append({'type': 'table', 'rows': table_rows, 'alignments': table_alignments})
            i += consumed_count
            continue

        # 普通段落
        else:
            paragraph_lines.append(stripped)

        i += 1

    if in_code_block and code_content:
        elements.append({'type': 'code', 'content': '\n'.join(code_content), 'lang': code_lang})
    flush_paragraph()

    if footnotes:
        elements.append({'type': 'footnotes', 'items': footnotes})

    return elements


def extract_document_title(elements: List[dict]) -> str:
    """提取文档标题，优先使用 Markdown 一级标题。"""
    for elem in elements:
        if elem.get('type') == 'title' and elem.get('content'):
            return str(elem.get('content'))

    for elem in elements:
        if elem.get('type') in {'heading2', 'heading3', 'heading4'} and elem.get('content'):
            return str(elem.get('content'))

    for elem in elements:
        if elem.get('type') == 'body' and elem.get('content'):
            content = str(elem.get('content')).strip()
            if content:
                return content[:48]

    return '无标题文档'


# ─────────────────────────────────────────────
# DOCX 元素创建
# ─────────────────────────────────────────────

def append_hyperlink_run(para, text: str, url: str, font_cfg: dict) -> None:
    """添加外部超链接 run。"""
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'Hyperlink')
    rPr.append(r_style)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    font_name = font_cfg.get('name', '宋体')
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)
    r_fonts.set(qn('w:cs'), font_name)
    rPr.append(r_fonts)

    size = OxmlElement('w:sz')
    size.set(qn('w:val'), str(int(font_cfg.get('size', 12) * 2)))
    rPr.append(size)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        text_node.set(qn('xml:space'), 'preserve')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


def add_inline_runs(para, content: str, font_cfg: dict) -> None:
    """根据 parse_inline 结果向段落追加 run。"""
    for text, attrs in parse_inline(content):
        if not text:
            continue

        link = attrs.get('link')
        if link and urlparse(link).scheme in {'http', 'https', 'mailto'}:
            append_hyperlink_run(para, text, link, font_cfg)
            continue

        run = para.add_run(text)
        font_name = 'Consolas' if attrs.get('code') else font_cfg.get('name', '宋体')
        font_size = 10 if attrs.get('code') else font_cfg.get('size', 12)
        is_bold = attrs.get('bold', font_cfg.get('bold', False))
        is_italic = attrs.get('italic', font_cfg.get('italic', False))
        set_run_font(run, font_name, font_size, bold=is_bold, italic=is_italic)

        if attrs.get('strike'):
            run.font.strike = True
        if attrs.get('superscript'):
            run.font.superscript = True


def apply_quote_decoration(para, fill: str = 'F5F5F5', border_color: str = '888888') -> None:
    """添加引用/callout 的左边框与背景。"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def set_outline_level(para, level: int) -> None:
    """写入 Word 大纲级别，让导航窗格和 TOC 能识别标题。"""
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(max(0, min(level, 8))))
    pPr.append(outline)


def apply_native_heading_style(para, style_key: str) -> None:
    """把视觉标题升级为 Word 原生标题结构。"""
    style_map = {
        'title': ('Title', None),
        'heading1': ('Heading 1', 0),
        'heading2': ('Heading 2', 1),
        'heading3': ('Heading 3', 2),
    }
    style_name, outline_level = style_map.get(style_key, ('Normal', None))
    try:
        para.style = style_name
    except KeyError:
        pass

    if outline_level is not None:
        set_outline_level(para, outline_level)


def add_paragraph_with_style(doc: Document, content: str, style_key: str, rules: dict) -> None:
    """添加带样式的段落"""
    style_cfg = rules.get(style_key, rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = style_cfg.get('paragraph', {})

    para = doc.add_paragraph()
    if style_key in ('title', 'heading1', 'heading2', 'heading3'):
        apply_native_heading_style(para, style_key)
    set_paragraph_format(para, para_cfg)

    # 标题：与下段保持同页，防止孤立标题
    if style_key in ('title', 'heading1', 'heading2', 'heading3'):
        para.paragraph_format.keep_with_next = True

    add_inline_runs(para, content, font_cfg)


def add_quote_paragraph(doc: Document, content: str, rules: dict) -> None:
    """添加引用段落（带左边框）"""
    style_cfg = rules.get('quote', rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = style_cfg.get('paragraph', {})

    para = doc.add_paragraph()
    set_paragraph_format(para, para_cfg)

    add_inline_runs(para, content, font_cfg)
    apply_quote_decoration(para)


def add_callout_paragraph(doc: Document, title: str, content: str, callout_type: str, rules: dict) -> None:
    """添加 Obsidian callout，转为正式 Word 提示块。"""
    style_cfg = rules.get('quote', rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = style_cfg.get('paragraph', {})

    color_map = {
        'warning': ('FFF2CC', 'D6B656'),
        'danger': ('F8CECC', 'B85450'),
        'failure': ('F8CECC', 'B85450'),
        'success': ('D5E8D4', '82B366'),
        'tip': ('D5E8D4', '82B366'),
        'question': ('DAE8FC', '6C8EBF'),
        'info': ('DAE8FC', '6C8EBF'),
        'note': ('F5F5F5', '888888'),
    }
    fill, border = color_map.get(callout_type, ('F5F5F5', '888888'))

    para = doc.add_paragraph()
    set_paragraph_format(para, para_cfg)
    para.paragraph_format.keep_together = True

    label_run = para.add_run(f"{title}：")
    set_run_font(label_run, font_cfg.get('name', '楷体'), font_cfg.get('size', 11),
                 bold=True, italic=False)
    if content:
        add_inline_runs(para, content, font_cfg)

    apply_quote_decoration(para, fill=fill, border_color=border)


def _ensure_numbering_definitions(doc: Document, max_level: int = 3) -> str:
    """确保文档有中文正式编号的 abstractNum 定义，返回 numId。

    定义三级编号体系：
      一级: 一、二、三（chineseCountingThousand + "、"）
      二级: （一）（二）（三）（中文括号 + chineseCounting）
      三级: 1. 2. 3.（decimal + "."）
    """
    # 检查是否已经创建过
    numbering_part = doc.part.numbering_part
    if numbering_part is not None:
        numbering_el = numbering_part._element
        # 简单方式：查找有没有我们定义的 abstractNum（用特定标识）
        for abnum in numbering_el.findall(qn('w:abstractNum')):
            for lvl in abnum.findall(qn('w:lvl')):
                start = lvl.find(qn('w:start'))
                if start is not None and start.get(qn('w:val')) == '1':
                    lvl_fmt = lvl.find(qn('w:numFmt'))
                    if lvl_fmt is not None and lvl_fmt.get(qn('w:val')) == 'chineseCountingThousand':
                        # 已存在，找对应的 numId
                        ab_id = abnum.get(qn('w:abstractNumId'))
                        for num in numbering_el.findall(qn('w:num')):
                            if num.get(qn('w:abstractNumId')) == ab_id:
                                return num.get(qn('w:numId'))

    # 定义 abstractNum XML
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    nsmap = {'w': W}

    # 确定下一个 abstractNumId
    next_ab_id = 0
    if numbering_part is not None:
        existing = numbering_part._element.findall(qn('w:abstractNum'))
        if existing:
            next_ab_id = max(int(e.get(qn('w:abstractNumId'), '0')) for e in existing) + 1

    abnum = OxmlElement('w:abstractNum')
    abnum.set(qn('w:abstractNumId'), str(next_ab_id))

    # 多级编号样式定义
    level_configs = [
        {
            'ilvl': 0, 'start': 1, 'numFmt': 'chineseCountingThousand',
            'lvlText': '%1、', 'lvlJc': 'left',
            'ind': 0, 'hanging': 420, 'font': '黑体',
        },
        {
            'ilvl': 1, 'start': 1, 'numFmt': 'chineseCounting',
            'lvlText': '（%2）', 'lvlJc': 'left',
            'ind': 420, 'hanging': 420, 'font': '楷体',
        },
        {
            'ilvl': 2, 'start': 1, 'numFmt': 'decimal',
            'lvlText': '%3.', 'lvlJc': 'left',
            'ind': 840, 'hanging': 420, 'font': '宋体',
        },
    ]

    for cfg in level_configs:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(cfg['ilvl']))

        start_el = OxmlElement('w:start')
        start_el.set(qn('w:val'), str(cfg['start']))
        lvl.append(start_el)

        fmt_el = OxmlElement('w:numFmt')
        fmt_el.set(qn('w:val'), cfg['numFmt'])
        lvl.append(fmt_el)

        txt_el = OxmlElement('w:lvlText')
        txt_el.set(qn('w:val'), cfg['lvlText'])
        lvl.append(txt_el)

        jc_el = OxmlElement('w:lvlJc')
        jc_el.set(qn('w:val'), cfg['lvlJc'])
        lvl.append(jc_el)

        # 段落缩进
        ppr = OxmlElement('w:pPr')
        ind_el = OxmlElement('w:ind')
        ind_el.set(qn('w:left'), str(cfg['ind']))
        ind_el.set(qn('w:hanging'), str(cfg['hanging']))
        ppr.append(ind_el)
        lvl.append(ppr)

        # 编号字体
        rpr = OxmlElement('w:rPr')
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:hint'), 'default')
        rpr.append(rfonts)
        lvl.append(rpr)

        abnum.append(lvl)

    # 无序列表 abstractNum（下一 ID）
    bullet_ab_id = next_ab_id + 1
    bullet_abnum = OxmlElement('w:abstractNum')
    bullet_abnum.set(qn('w:abstractNumId'), str(bullet_ab_id))

    bullet_levels = [
        {'ilvl': 0, 'indent': 0},
        {'ilvl': 1, 'indent': 420},
        {'ilvl': 2, 'indent': 840},
    ]
    bullet_chars = ['•', '◦', '▪']

    for cfg, char in zip(bullet_levels, bullet_chars):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(cfg['ilvl']))

        start_el = OxmlElement('w:start')
        start_el.set(qn('w:val'), '1')
        lvl.append(start_el)

        fmt_el = OxmlElement('w:numFmt')
        fmt_el.set(qn('w:val'), 'bullet')
        lvl.append(fmt_el)

        txt_el = OxmlElement('w:lvlText')
        txt_el.set(qn('w:val'), char)
        lvl.append(txt_el)

        jc_el = OxmlElement('w:lvlJc')
        jc_el.set(qn('w:val'), 'left')
        lvl.append(jc_el)

        ppr = OxmlElement('w:pPr')
        ind_el = OxmlElement('w:ind')
        ind_el.set(qn('w:left'), str(cfg['indent'] + 420))
        ind_el.set(qn('w:hanging'), '420')
        ppr.append(ind_el)
        lvl.append(ppr)

        rpr = OxmlElement('w:rPr')
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), 'Symbol')
        rfonts.set(qn('w:hAnsi'), 'Symbol')
        rfonts.set(qn('w:hint'), 'default')
        rpr.append(rfonts)
        lvl.append(rpr)

        bullet_abnum.append(lvl)

    # 添加到 numbering part
    if numbering_part is None:
        from docx.opc.part import Part
        num_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '</w:numbering>'
        )
        num_part = Part(
            PackURI('/word/numbering.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            num_xml.encode('utf-8'),
            doc.part.package
        )
        doc.part.relate_to(num_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
        numbering_part = num_part

    numbering_el = numbering_part._element if hasattr(numbering_part, '_element') else ET.fromstring(numbering_part.blob.decode('utf-8'))
    if not hasattr(numbering_part, '_element'):
        numbering_part._element = numbering_el

    # 插入 abstractNum（必须在 num 之前）
    existing_nums = numbering_el.findall(qn('w:num'))
    if existing_nums:
        first_num = existing_nums[0]
        first_num.addprevious(abnum)
        first_num.addprevious(bullet_abnum)
    else:
        numbering_el.append(abnum)
        numbering_el.append(bullet_abnum)

    # 创建 num 元素关联 abstractNum
    next_num_id = 1
    existing_num_ids = numbering_el.findall(qn('w:num'))
    if existing_num_ids:
        next_num_id = max(int(e.get(qn('w:numId'), '0')) for e in existing_num_ids) + 1

    ordered_num = OxmlElement('w:num')
    ordered_num.set(qn('w:numId'), str(next_num_id))
    ordered_abref = OxmlElement('w:abstractNumId')
    ordered_abref.set(qn('w:val'), str(next_ab_id))
    ordered_num.append(ordered_abref)
    numbering_el.append(ordered_num)

    bullet_num_id = next_num_id + 1
    bullet_num = OxmlElement('w:num')
    bullet_num.set(qn('w:numId'), str(bullet_num_id))
    bullet_abref = OxmlElement('w:abstractNumId')
    bullet_abref.set(qn('w:val'), str(bullet_ab_id))
    bullet_num.append(bullet_abref)
    numbering_el.append(bullet_num)

    # 保存回 part
    if hasattr(numbering_part, 'blob'):
        numbering_part._blob = ET.tostring(numbering_el, encoding='unicode', xml_declaration=True).encode('utf-8')

    # 缓存到 document 对象上，后续 add_list_item 使用
    doc._docforge_num_ids = {
        'ordered': str(next_num_id),
        'bullet': str(bullet_num_id),
    }
    return str(next_num_id)


def apply_list_style(para, ordered: bool, level: int, doc: Document) -> None:
    """使用显式 w:numPr 写入编号属性，支持中文正式文档多级编号。

    有序：一、（一）1.
    无序：• ◦ ▪
    """
    safe_level = max(0, min(level, 2))

    num_ids = getattr(doc, '_docforge_num_ids', None)
    if not num_ids:
        # 编号定义创建失败时 fallback：在文本前手动插入编号符号
        prefix = ''
        if ordered:
            chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
            if safe_level == 0:
                prefix = '一、'
            elif safe_level == 1:
                prefix = '（一）'
            else:
                prefix = '1. '
        else:
            prefix = '• ' if safe_level == 0 else ('◦ ' if safe_level == 1 else '▪ ')

        # 在段落最前面插入编号前缀
        if prefix:
            existing = para.text
            para.clear()
            prefix_run = para.add_run(prefix)
            set_run_font(prefix_run, font_cfg.get('name', '宋体') if font_cfg else '宋体', font_cfg.get('size', 12) if font_cfg else 12, bold=True)
            if existing:
                para.add_run(existing)

        para.paragraph_format.left_indent = Inches(0.35 + safe_level * 0.35)
        para.paragraph_format.first_line_indent = Inches(-0.3)
        return

    num_id = num_ids['ordered'] if ordered else num_ids['bullet']

    pPr = para._p.get_or_add_pPr()

    # 移除已有的 numPr（如果重新设置）
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(safe_level))
    numPr.append(ilvl)

    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), num_id)
    numPr.append(numId_el)

    # ind 确保在 numPr 后面
    pPr.insert(0, numPr)

    para.paragraph_format.left_indent = Inches(0.35 + safe_level * 0.35)
    para.paragraph_format.first_line_indent = Inches(-0.3)


def add_list_item(doc: Document, elem: dict, rules: dict) -> None:
    """添加列表项，支持 Obsidian 任务列表状态和中文正式编号。"""
    style_cfg = rules.get('list', rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = dict(style_cfg.get('paragraph', {}))
    para_cfg['indent_first_line'] = 0

    para = doc.add_paragraph()
    set_paragraph_format(para, para_cfg)
    apply_list_style(para, bool(elem.get('ordered')), int(elem.get('level') or 0), doc)

    task_state = elem.get('task_state')
    if task_state == 'done':
        prefix_run = para.add_run('完成：')
        set_run_font(prefix_run, font_cfg.get('name', '宋体'), font_cfg.get('size', 12), bold=True)
    elif task_state == 'todo':
        prefix_run = para.add_run('待办：')
        set_run_font(prefix_run, font_cfg.get('name', '宋体'), font_cfg.get('size', 12), bold=True)

    add_inline_runs(para, elem.get('content', ''), font_cfg)


def add_code_paragraph(doc: Document, content: str, rules: dict, lang: str = '') -> None:
    """添加代码块段落"""
    style_cfg = rules.get('code', rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = style_cfg.get('paragraph', {})

    if lang:
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(6)
        label.paragraph_format.space_after = Pt(0)
        run = label.add_run(f"代码：{lang}")
        set_run_font(run, font_cfg.get('name', 'Consolas'), max(font_cfg.get('size', 10) - 1, 8), bold=True)

    for line in content.split('\n'):
        para = doc.add_paragraph()
        set_paragraph_format(para, para_cfg)

        run = para.add_run(line if line else ' ')
        set_run_font(run, font_cfg.get('name', 'Consolas'), font_cfg.get('size', 10))

        # 灰色背景
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)


def resolve_asset_path(src: str, asset_root: Optional[str]) -> Optional[str]:
    """解析 Markdown/Obsidian 图片路径。"""
    if not src:
        return None

    parsed = urlparse(src)
    if parsed.scheme in {'http', 'https'}:
        return None

    cleaned = unquote(src.split('#', 1)[0].split('|', 1)[0]).strip()
    if cleaned.startswith('<') and cleaned.endswith('>'):
        cleaned = cleaned[1:-1].strip()

    allowed_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.gif'}
    root = Path(asset_root).resolve() if asset_root else None
    candidates: List[Path] = []
    path_candidate = Path(cleaned)
    if path_candidate.is_absolute() and root is None:
        candidates.append(path_candidate)
    if root:
        candidates.append(root / cleaned)
        candidates.append(root / 'attachments' / cleaned)
        candidates.append(root / 'assets' / cleaned)
        candidates.append(root / 'images' / cleaned)

    for candidate in candidates:
        try:
            resolved = candidate.resolve()
            if root and not (resolved == root or root in resolved.parents):
                continue
            if resolved.suffix.lower() not in allowed_extensions:
                continue
            if resolved.exists() and resolved.is_file() and resolved.stat().st_size <= 20 * 1024 * 1024:
                return str(resolved)
        except OSError:
            continue
    return None


def add_image(doc: Document, src: str, alt: str, rules: dict, asset_root: Optional[str]) -> None:
    """添加图片；找不到本地图片时保留清晰占位说明，避免静默丢内容。"""
    image_path = resolve_asset_path(src, asset_root)
    caption = alt or Path(src).name

    if not image_path:
        add_quote_paragraph(doc, f"图片未嵌入：{caption}（{src}）", rules)
        return

    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(5.8))

        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            set_run_font(caption_run, '宋体', 9, italic=True, color_hex='666666')
            caption_para.paragraph_format.space_after = Pt(6)
    except Exception as exc:
        add_quote_paragraph(doc, f"图片嵌入失败：{caption}（{exc}）", rules)


def _parse_table_rows(lines: List[str], start_i: int) -> Tuple[List[List[str]], List[str], int]:
    """解析 Markdown 表格行，支持对齐行、escaped pipe (\\|) 和列数归一化。

    返回 (rows, alignments, consumed_count)，consumed_count 含对齐行。
    """
    rows = []
    alignments = []
    raw_lines = []
    i = start_i

    while i < len(lines):
        row_line = lines[i].strip()
        if '|' not in row_line or not row_line.startswith('|'):
            break
        raw_lines.append(row_line)
        i += 1

    if not raw_lines:
        return rows, alignments

    # 检测对齐行（第二行，格式如 |:---|:---:|---:|）
    if len(raw_lines) >= 2:
        second = raw_lines[1]
        if re.match(r'^\|[\s\-:]+\|', second):
            align_cells = [c.strip() for c in second.strip('|').split('|')]
            for cell in align_cells:
                cell = cell.strip()
                if cell.startswith(':') and cell.endswith(':'):
                    alignments.append('center')
                elif cell.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')
            raw_lines.pop(1)  # 移除对齐行

    # 解析单元格，处理 escaped pipe
    def split_row(line: str) -> List[str]:
        # 先把 \| 替换为占位符，split 后再恢复
        placeholder = '\x00PIPE\x00'
        line = line.replace('\\|', placeholder)
        # 精确去掉首尾各一个 |，而非 strip（避免连续 || 空单元格被误剥）
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = [c.strip().replace(placeholder, '|') for c in line.split('|')]
        return cells

    for line in raw_lines:
        cells = split_row(line)
        rows.append(cells)

    # 归一化列数
    if rows:
        num_cols = max(len(row) for row in rows)
        for row in rows:
            while len(row) < num_cols:
                row.append('')
        while len(alignments) < num_cols:
            alignments.append('left')

    return rows, alignments, len(raw_lines)


def add_table(doc: Document, rows: List[List[str]], rules: dict, alignments: Optional[List[str]] = None) -> None:
    """添加表格，根据页面边距计算列宽，支持对齐和表头重复。"""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False

    body_cfg = rules.get('body', {})
    font_cfg = body_cfg.get('font', {})

    # 根据页面边距计算表格宽度
    try:
        section = doc.sections[-1]
        page_width = Inches(8.27)  # A4 默认
        if section.page_width:
            page_width = section.page_width
        left_margin = section.left_margin or Inches(1.25)
        right_margin = section.right_margin or Inches(1.25)
        available = page_width - left_margin - right_margin
        column_width = available / max(num_cols, 1)
    except Exception:
        column_width = Inches(6.2 / max(num_cols, 1))

    if alignments is None:
        alignments = ['left'] * num_cols

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        is_header = (i == 0)

        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = ''
            cell.width = column_width
            para = cell.paragraphs[0]

            # 表头居中，正文按对齐设置
            align = WD_ALIGN_PARAGRAPH.CENTER if is_header else {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(alignments[j] if j < len(alignments) else 'left', WD_ALIGN_PARAGRAPH.LEFT)
            para.alignment = align

            if is_header:
                run = para.add_run(cell_text)
                set_run_font(run, font_cfg.get('name', '宋体'), font_cfg.get('size', 12), bold=True)
            else:
                add_inline_runs(para, cell_text, font_cfg)

            # 表头背景 + 重复标题行
            if is_header:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DAEEF3')
                tcPr.append(shd)

                # 重复标题行（跨页时表头自动重复）
                tblHeader = OxmlElement('w:tblHeader')
                tcPr.append(tblHeader)

    # 写入 tblGrid 确保列宽生效
    try:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            tbl.remove(tblGrid)
        tblGrid = OxmlElement('w:tblGrid')

        for _ in range(num_cols):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(column_width)))
            tblGrid.append(gridCol)

        # tblGrid 放在 tblPr 后面
        tblPr.addnext(tblGrid)
    except Exception:
        pass


def add_footnotes_section(doc: Document, items: List[dict], rules: dict) -> None:
    """将 Obsidian 脚注转换为 Word 原生脚注（w:footnoteReference + footnotes.xml）。

    python-docx 不直接支持脚注，需要操作底层 XML：
    1. 创建 word/footnotes.xml 并关联到文档包
    2. 在正文段落中插入 w:footnoteReference 上标引用
    3. 脚注编号从 1 开始自动递增
    """
    if not items:
        return

    try:
        _add_native_footnotes(doc, items, rules)
    except Exception as exc:
        # 原生脚注创建失败时 fallback 为文末注，确保内容不丢失
        print(f'警告: 原生脚注创建失败，回退为文末注: {exc}', file=sys.stderr)
        _add_endnote_fallback(doc, items, rules)


def _add_endnote_fallback(doc: Document, items: List[dict], rules: dict) -> None:
    """Fallback: 将脚注以文末注形式呈现，确保内容不丢失。"""
    add_paragraph_with_style(doc, '注释', 'heading1', rules)
    for item in items:
        content = f"[{item.get('id')}] {item.get('content', '')}"
        style_cfg = rules.get('body', {})
        font_cfg = style_cfg.get('font', {})
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.first_line_indent = Inches(-0.2)
        add_inline_runs(para, content, font_cfg)


def _add_native_footnotes(doc: Document, items: List[dict], rules: dict) -> None:
    """创建 Word 原生脚注。仅在 try 块内调用。"""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 直接构建 XML 字符串，避免 ElementTree 命名空间问题
    xml_parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        # 分隔符脚注（ID 0 是 Word 保留的）
        '<w:footnote w:type="separator" w:id="0"><w:p><w:r><w:sep/></w:r></w:p></w:footnote>',
    ]

    footnote_ids: Dict[str, int] = {}
    for idx, item in enumerate(items, start=1):
        footnote_ids[item.get('id', '')] = idx
        content = item.get('content', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        xml_parts.append(
            f'<w:footnote w:id="{idx}">'
            f'<w:p>'
            f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:t>{idx}</w:t></w:r>'
            f'<w:r><w:t xml:space="preserve"> </w:t></w:r>'
            f'<w:r><w:t xml:space="preserve">{content}</w:t></w:r>'
            f'</w:p>'
            f'</w:footnote>'
        )

    xml_parts.append('</w:footnotes>')
    footnotes_xml = ''.join(xml_parts)

    from docx.opc.part import Part

    footnotes_part = Part(
        PackURI('/word/footnotes.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
        footnotes_xml.encode('utf-8'),
        doc.part.package
    )
    doc.part.relate_to(footnotes_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')

    # 在正文中找到脚注引用（上标 [id]）并替换为 w:footnoteReference
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.superscript:
                continue
            text = run.text
            if not text:
                continue
            fn_match = re.match(r'^\[(.+)\]$', text)
            if fn_match and fn_match.group(1) in footnote_ids:
                fn_id = footnote_ids[fn_match.group(1)]
                run.text = ''
                r_elem = run._r
                rpr = r_elem.find(qn('w:rPr'))
                if rpr is not None:
                    vert = rpr.find(qn('w:vertAlign'))
                    if vert is not None:
                        rpr.remove(vert)
                fn_ref = OxmlElement('w:footnoteReference')
                fn_ref.set(qn('w:id'), str(fn_id))
                r_elem.append(fn_ref)


def enable_update_fields_on_open(doc: Document) -> None:
    """让 Word 打开文档时主动更新目录、页码等域。"""
    try:
        settings = doc.settings._element
        existing = settings.find(qn('w:updateFields'))
        if existing is None:
            existing = OxmlElement('w:updateFields')
            settings.append(existing)
        existing.set(qn('w:val'), 'true')
    except Exception:
        pass


def add_table_of_contents(doc: Document, rules: dict) -> None:
    """插入可由 Word 更新的目录域。"""
    style_cfg = rules.get('heading1', rules['body'])
    font_cfg = style_cfg.get('font', {})
    para_cfg = style_cfg.get('paragraph', {})

    title_para = doc.add_paragraph()
    set_paragraph_format(title_para, para_cfg)
    title_para.paragraph_format.keep_with_next = True
    title_run = title_para.add_run('目录')
    set_run_font(
        title_run,
        font_cfg.get('name', '黑体'),
        font_cfg.get('size', 16),
        bold=font_cfg.get('bold', True),
        italic=font_cfg.get('italic', False)
    )

    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_before = Pt(0)
    toc_para.paragraph_format.space_after = Pt(12)
    run = toc_para.add_run()
    add_field(run, r'TOC \o "1-3" \h \z \u', '请在 Word 中更新域以生成目录')


# ─────────────────────────────────────────────
# 主要功能
# ─────────────────────────────────────────────

def extract_styles_from_docx(docx_path: str) -> dict:
    """从 DOCX 模板提取样式规则，包括页面大小/方向、页眉页脚间距等。"""
    rules = default_style_rules()

    try:
        doc = Document(docx_path)

        # 提取正文默认样式
        try:
            normal = doc.styles['Normal']
            if normal.font.name:
                rules['body']['font']['name'] = normal.font.name
            if normal.font.size:
                rules['body']['font']['size'] = round(normal.font.size.pt)
        except Exception:
            pass

        # 提取标题样式
        heading_map = {
            'Heading 1': 'heading1',
            'Heading 2': 'heading2',
            'Heading 3': 'heading3',
        }
        for word_style, key in heading_map.items():
            try:
                style = doc.styles[word_style]
                if style.font.name:
                    rules[key]['font']['name'] = style.font.name
                if style.font.size:
                    rules[key]['font']['size'] = round(style.font.size.pt)
                if style.font.bold is not None:
                    rules[key]['font']['bold'] = style.font.bold
            except Exception:
                pass

        # 提取页面设置：边距、页面大小、方向
        try:
            section = doc.sections[0]
            rules['page_margin']['top'] = round(section.top_margin.inches, 2)
            rules['page_margin']['bottom'] = round(section.bottom_margin.inches, 2)
            rules['page_margin']['left'] = round(section.left_margin.inches, 2)
            rules['page_margin']['right'] = round(section.right_margin.inches, 2)

            # 页面大小和方向
            page_w = section.page_width
            page_h = section.page_height
            if page_w and page_h:
                rules['page_size']['width'] = round(page_w.inches, 2)
                rules['page_size']['height'] = round(page_h.inches, 2)
                # 判断方向：宽 > 高为 landscape
                if page_w > page_h:
                    rules['page_size']['orientation'] = 'landscape'
                else:
                    rules['page_size']['orientation'] = 'portrait'

            # 页眉页脚距离
            try:
                rules['page_margin']['header_distance'] = round(section.header_distance.inches, 2)
                rules['page_margin']['footer_distance'] = round(section.footer_distance.inches, 2)
            except Exception:
                pass

            # 装订线
            try:
                gutter = section.gutter
                if gutter and gutter.inches > 0:
                    rules['page_margin']['gutter'] = round(gutter.inches, 2)
            except Exception:
                pass

        except Exception:
            pass

        # 许多中文正式文档不用 Word Heading 样式，而是直接给段落设置字体字号。
        # 因此再从实际段落中抽样提取，覆盖默认样式。
        paragraphs = [p for p in doc.paragraphs if normalized_para_text(p)]
        if paragraphs:
            title_para = paragraphs[0]
            rules['title'] = paragraph_to_rule(title_para, rules['title'])

        body_para = next((p for p in paragraphs if len(normalized_para_text(p)) >= 50), None)
        if body_para is not None:
            rules['body'] = paragraph_to_rule(body_para, rules['body'])

        seen_headings = set()
        for para in paragraphs[:120]:
            text = normalized_para_text(para)
            level = heading_level_from_text(text)
            if not level or level in seen_headings:
                continue
            rules[level] = paragraph_to_rule(para, rules[level])
            seen_headings.add(level)
            if len(seen_headings) == 3:
                break

    except Exception as e:
        print(f"警告: 无法解析模板样式，使用默认值: {e}", file=sys.stderr)

    return rules


def markdown_to_docx(
    markdown: str,
    output_path: str,
    style_rules: Optional[dict] = None,
    asset_root: Optional[str] = None
) -> str:
    """将 Markdown 转换为带样式的 DOCX"""
    rules = style_rules if style_rules else default_style_rules()

    doc = Document()

    # 设置页面：边距、大小、方向
    margin = rules.get('page_margin', {})
    page_size = rules.get('page_size', {})
    section = doc.sections[0]
    section.top_margin = Inches(margin.get('top', 1.0))
    section.bottom_margin = Inches(margin.get('bottom', 1.0))
    section.left_margin = Inches(margin.get('left', 1.25))
    section.right_margin = Inches(margin.get('right', 1.25))

    # 页面大小和方向
    if page_size.get('width') and page_size.get('height'):
        section.page_width = Inches(page_size['width'])
        section.page_height = Inches(page_size['height'])
        if page_size.get('orientation') == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
    # 页眉页脚距离
    if margin.get('header_distance'):
        section.header_distance = Inches(margin['header_distance'])
    if margin.get('footer_distance'):
        section.footer_distance = Inches(margin['footer_distance'])
    # 装订线
    if margin.get('gutter'):
        section.gutter = Inches(margin['gutter'])

    # 解析 Markdown
    elements = parse_markdown(markdown)
    doc.core_properties.title = extract_document_title(elements)
    doc.core_properties.author = 'DocForge'
    doc.core_properties.subject = 'DocForge 自动生成文档'
    doc.core_properties.comments = 'Generated by DocForge'
    should_add_toc = sum(1 for elem in elements if elem.get('type') in {'heading2', 'heading3', 'heading4'}) >= 2
    toc_inserted = False

    # 如果有列表元素，初始化中文多级编号定义
    has_lists = any(e.get('type') == 'list_item' for e in elements)
    if has_lists:
        try:
            _ensure_numbering_definitions(doc)
        except Exception:
            pass  # 编号定义失败不阻塞文档生成

    for elem in elements:
        etype = elem['type']

        if etype == 'title':
            add_paragraph_with_style(doc, elem['content'], 'title', rules)
            if should_add_toc and not toc_inserted:
                add_table_of_contents(doc, rules)
                toc_inserted = True

        elif etype == 'heading2':
            if should_add_toc and not toc_inserted:
                add_table_of_contents(doc, rules)
                toc_inserted = True
            add_paragraph_with_style(doc, elem['content'], 'heading1', rules)

        elif etype == 'heading3':
            add_paragraph_with_style(doc, elem['content'], 'heading2', rules)

        elif etype == 'heading4':
            add_paragraph_with_style(doc, elem['content'], 'heading3', rules)

        elif etype == 'body':
            add_paragraph_with_style(doc, elem['content'], 'body', rules)

        elif etype == 'list_item':
            add_list_item(doc, elem, rules)

        elif etype == 'quote':
            add_quote_paragraph(doc, elem['content'], rules)

        elif etype == 'callout':
            add_callout_paragraph(
                doc,
                elem.get('title', ''),
                elem.get('content', ''),
                elem.get('callout_type', ''),
                rules
            )

        elif etype == 'code':
            add_code_paragraph(doc, elem['content'], rules, elem.get('lang', ''))

        elif etype == 'image':
            add_image(doc, elem.get('src', ''), elem.get('alt', ''), rules, asset_root)

        elif etype == 'table':
            add_table(doc, elem['rows'], rules, elem.get('alignments'))

        elif etype == 'footnotes':
            add_footnotes_section(doc, elem.get('items', []), rules)

        elif etype == 'hr':
            para = doc.add_paragraph()
            run = para.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    add_page_number_footer(doc)
    enable_update_fields_on_open(doc)
    doc.save(output_path)
    errors = validate_docx_package(output_path)
    if errors:
        raise RuntimeError('DOCX 校验失败: ' + '; '.join(errors))
    return output_path


# ─────────────────────────────────────────────
# 命令行入口
# ─────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("用法:", file=sys.stderr)
        print("  提取样式: python docforge_py.py extract <input.docx> <output.json>", file=sys.stderr)
        print("  生成文档: python docforge_py.py generate <input.md> <output.docx> [--style <style.json>] [--asset-root <dir>]", file=sys.stderr)
        print("  校验文档: python docforge_py.py validate <input.docx>", file=sys.stderr)
        sys.exit(1)

    command = sys.argv[1]

    if command == "extract":
        if len(sys.argv) < 4:
            print("错误: 需要指定输出文件", file=sys.stderr)
            sys.exit(1)
        docx_path = sys.argv[2]
        output_path = sys.argv[3]

        if not os.path.exists(docx_path):
            print(f"错误: 文件不存在: {docx_path}", file=sys.stderr)
            sys.exit(1)

        styles = extract_styles_from_docx(docx_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(styles, f, ensure_ascii=False, indent=2)
        print(f"样式已提取到: {output_path}")

    elif command == "generate":
        if len(sys.argv) < 4:
            print("错误: 需要指定输出文件", file=sys.stderr)
            sys.exit(1)

        md_path = sys.argv[2]
        output_path = sys.argv[3]

        style_file = None
        asset_root = str(Path(md_path).resolve().parent)
        for i in range(4, len(sys.argv)):
            if sys.argv[i] == "--style" and i + 1 < len(sys.argv):
                style_file = sys.argv[i + 1]
            if sys.argv[i] == "--asset-root" and i + 1 < len(sys.argv):
                asset_root = sys.argv[i + 1]

        if not os.path.exists(md_path):
            print(f"错误: 文件不存在: {md_path}", file=sys.stderr)
            sys.exit(1)

        with open(md_path, 'r', encoding='utf-8') as f:
            markdown = f.read()

        style_rules = None
        if style_file and os.path.exists(style_file):
            with open(style_file, 'r', encoding='utf-8') as f:
                style_rules = json.load(f)

        result_path = markdown_to_docx(markdown, output_path, style_rules, asset_root)
        print(f"DOCX 已生成: {result_path}")

    elif command == "validate":
        docx_path = sys.argv[2]
        errors = validate_docx_package(docx_path)
        if errors:
            print(json.dumps({"valid": False, "errors": errors}, ensure_ascii=False), file=sys.stderr)
            sys.exit(1)
        print(json.dumps({"valid": True, "errors": []}, ensure_ascii=False))

    else:
        print(f"未知命令: {command}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
